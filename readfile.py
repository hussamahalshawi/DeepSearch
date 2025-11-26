import json
import os
import pandas as pd
import yaml
import xml.etree.ElementTree as ET
from docx import Document
from PyPDF2 import PdfReader
from PIL import Image
import wave
import cv2
import zipfile
from bs4 import BeautifulSoup
import sqlite3
from zipfile import ZipFile
from io import BytesIO
from transformers import pipeline
from pygments.lexer import words



def get_urls_files(directory):
    urls = {}
    for root, dirs, files in os.walk(directory):
        if ".git" not in  root:
            for file in files:
                urls[file] = os.path.join(root, file)
    return urls


class Read:
    def __init__(self, directory: str):
        """
        Read url file

        Args:
            directory (str): url file.
        """
        try:
            self.directory = directory
            self.urls = get_urls_files(directory)
            self.data_all = {}
            self.split_name_file(self.urls)
            # self.read_file(self.urls)
        except Exception as e:
            print(f'Error: {e}')

    def split_name_file(self, urls: dict, z = None, urlz=None) -> None:
        urlzn = ""
        if urlz is None:
            urlz = []
        for index, (name_file, url) in enumerate(urls.items(), start=0):
            self.data_all[url] = []
            lis_name_file = name_file.split(".")
            if urlz != []:
                urlzn = urlz[index]
                self.data_all[urlzn] = []
                self.data_all[urlzn].append(lis_name_file[0])
            else:
                self.data_all[url].append(lis_name_file[0])
            print(lis_name_file[-1])
            if lis_name_file[-1] == "txt":
                self.read_file_txt(url, z, urlzn)
            elif lis_name_file[-1] == "json":
                self.read_file_json(url, z, urlzn)
            elif lis_name_file[-1] == "csv":
                self.read_file_csv(url, z, urlzn)
            elif lis_name_file[-1] == "yaml" or lis_name_file[-1] == "yml":
                self.read_file_yaml(url, z, urlzn)
            elif lis_name_file[-1] == "xml":
                self.read_file_xml(url, z, urlzn)
            elif lis_name_file[-1] == "xlsx" or lis_name_file[-1] == "xls":
                self.read_file_xlsx(url, z, urlzn)
            elif lis_name_file[-1] == "docx" or lis_name_file[-1] == "doc":
                self.read_file_docx(url, z, urlzn)
            elif lis_name_file[-1] == "pdf":
                self.read_file_pdf(url, z, urlzn)
            # elif lis_name_file[-1] == "png":###########
            #     self.read_file_image(url, z, urlzn)
            # elif lis_name_file[-1] == "mp3":###########
            #     self.read_file_audio(url, z, urlzn)
            # elif lis_name_file[-1] == "mp4":  ###########
            #     self.read_file_video(url, z, urlzn)
            elif lis_name_file[-1] == "zip":  ###########
                self.read_file_zip(url, z, urlzn)
            elif lis_name_file[-1] == "html":  ###########
                self.read_file_html(url, z, urlzn)
            elif lis_name_file[-1] == "py" or lis_name_file[-1] == "md":  ###########
                self.read_file_txt(url, z, urlzn)
            elif lis_name_file[-1] == "db":  ###########
                self.read_file_databace(url, z, urlzn)
            else:
                self.read_file_txt(url, z, urlzn)

    def read_file_txt(self, url, z, urlz) -> None:
        try:
            if z:
                with z.open(url, "r") as file:
                    content = file.read()
                words = content.split()
                self.data_all[urlz].extend(words)
            else:
                with open(url, "r") as file:
                    content = file.read()
                words = content.split()
                self.data_all[url].extend(words)
        except FileNotFoundError:
            print("The file doesn't exist.")
        except Exception as e:
            print("An error occurred:", e)


    def read_file_json(self, url, z, urlz) -> None:
        try:
            if z:
                with z.open(url) as file:
                    content = json.load(file)
                self.extract_text_json_yaml(content, urlz)
            else:
                with open(url, "r") as file:
                    content = json.load(file)
                self.extract_text_json_yaml(content, url)

        except FileNotFoundError:
            print("The file doesn't exist.")
        except Exception as e:
            print("An error occurred:", e)

    def extract_text_json_yaml(self, data, url):
        try:
            if isinstance(data, dict):
                for key, value in data.items():
                    self.data_all[url].append(key)
                    self.extract_text_json_yaml(value, url)
            elif isinstance(data, list):
                for item in data:
                    self.extract_text_json_yaml(item, url)
            elif isinstance(data, int) or isinstance(data, float) or isinstance(data, str):
                self.data_all[url].append(data)

        except FileNotFoundError:
            print("The file doesn't exist.")
        except Exception as e:
            print("An error occurred:", e)


    def read_file_csv(self, url, z, urlz) -> None:
        try:
            df = pd.read_csv(url)
            if z:
                for col in df.columns:
                    self.data_all[urlz].append(col)
                    for value in df[col].astype(str):
                        for word in value.split():
                            self.data_all[urlz].append(word)
            else:
                for col in df.columns:
                    self.data_all[url].append(col)
                    for value in df[col].astype(str):
                        for word in value.split():
                            self.data_all[url].append(word)
        except FileNotFoundError:
            print("The file doesn't exist.")
        except Exception as e:
            print("An error occurred:", e)


    def read_file_yaml(self, url, z, urlz) -> None:
        try:
            if z:
                with z.open(url) as file:
                    content = yaml.safe_load(file)
                self.extract_text_json_yaml(content, urlz)
            else:
                with open(url) as file:
                    content = yaml.safe_load(file)
                self.extract_text_json_yaml(content, url)
        except FileNotFoundError:
            print("The file doesn't exist.")
        except Exception as e:
            print("An error occurred:", e)


    def read_file_xml(self, url, z, urlz) -> None:
        try:
            tree = ET.parse(url)
            root = tree.getroot()
            texts = [
                elem.text.strip()
                for elem in root.iter()
                if elem.text and elem.text.strip()
            ]
            words = []
            for item in texts:
                words.extend(item.split())
            if z:
                self.data_all[urlz].extend(words)
            else:
                self.data_all[url].extend(words)
        except FileNotFoundError:
            print("The file doesn't exist.")
        except Exception as e:
            print("An error occurred:", e)

    def read_file_xlsx(self, url, z, urlz) -> None:
        try:
            df = pd.read_excel(url)
            words = []
            for col in df.columns:
                if "Unnamed" not in str(col):
                    words.append(col)
                for value in df[col]:
                    if pd.notna(value):
                        value = str(value).strip()
                        if value:
                            words.extend(value.split())
            if z:
                self.data_all[urlz].extend(words)
            else:
                self.data_all[url].extend(words)
        except FileNotFoundError:
            print("The file doesn't exist.")
        except Exception as e:
            print("An error occurred:", e)


    def read_file_docx(self, url, z, urlz) -> None:
        try:
            doc = Document(url)
            if z:
                for p in doc.paragraphs:
                    self.data_all[urlz].extend(p.text.split())
            else:
                for p in doc.paragraphs:
                    self.data_all[url].extend(p.text.split())
        except FileNotFoundError:
            print("The file doesn't exist.")
        except Exception as e:
            print("An error occurred:", e)

    def read_file_pdf(self, url, z, urlz) -> None:
        try:
            reader = PdfReader(url)
            if z:
                for page in reader.pages:
                    self.data_all[urlz].extend(page.extract_text().split())
            else:
                for page in reader.pages:
                            self.data_all[url].extend(page.extract_text().split())
        except FileNotFoundError:
            print("The file doesn't exist.")
        except Exception as e:
            print("An error occurred:", e)


    def read_file_image(self, url, z, urlz) -> None:
        try:
            # الوصف بالعربي للصوره
            # captioner = pipeline(
            #     "image-to-text",
            #     model="microsoft/Florence-2-large",
            #     task="caption",
            #     max_new_tokens=70
            # )
            captioner = pipeline("image-to-text", model="Salesforce/blip-image-captioning-base")
            if z:
                # print(url)
                img = Image.open(BytesIO(z.read(url)))
                print("///",img)
            else:
                img = Image.open(url)
                print(img)
            result = captioner(img)
            description = result[0]["generated_text"]

            print("\n📌 وصف الصورة:")
            print(description)
        except FileNotFoundError:
            print("The file doesn't exist.")
        except Exception as e:
            print("An error occurred:", e)


    def read_file_audio(self, url, z, urlz) -> None:
        try:
            asr = pipeline("automatic-speech-recognition", model="openai/whisper-small")
            # asr = pipeline(
            #     "automatic-speech-recognition",
            #     model="openai/whisper-medium",
            #     generate_kwargs={"language": "arabic"}
            # )
            if z:  # لو الملف الصوتي داخل ZIP
                audio_bytes = BytesIO(z.read(url))  # url = مسار الملف داخل zip
                text = asr(audio_bytes)["text"]
            else:
                text = asr(url)["text"]

            print("\n📢 نص محتوى الصوت:")
            print(text)
            # audio = AudioSegment.from_file("file.mp3")
        except FileNotFoundError:
            print("The file doesn't exist.")
        except Exception as e:
            print("An error occurred:", e)

    def read_file_video(self, url, z, urlz) -> None:
        try:
            video = cv2.VideoCapture(url)

        except FileNotFoundError:
            print("The file doesn't exist.")
        except Exception as e:
            print("An error occurred:", e)

    def read_file_zip(self, url, z, urlz) -> None:
        try:
            urls = {}
            urlz = []
            with zipfile.ZipFile(url) as zip_obj:
                for file_info in zip_obj.infolist():
                    if file_info.is_dir():
                        continue
                    file_name = file_info.filename.split("/")[-1]
                    urls[file_name] = file_info.filename
                    urlz.append(url + "/" + file_info.filename)
                self.split_name_file(urls, zip_obj , urlz)
        except FileNotFoundError:
            print("The file doesn't exist.")
        except Exception as e:
            print("An error occurred:", e)

    def read_file_html(self, url, z, urlz) -> None:
        try:
            with open(url) as file:
                soup = BeautifulSoup(file, "html.parser")
        except FileNotFoundError:
            print("The file doesn't exist.")
        except Exception as e:
            print("An error occurred:", e)


    def read_file_code(self, url, z, urlz) -> None:
        try:
            with open(url) as file:
                content = file.read()
        except FileNotFoundError:
            print("The file doesn't exist.")
        except Exception as e:
            print("An error occurred:", e)


    def read_file_databace(self, url, z, urlz) -> None:
        try:
            conn = sqlite3.connect(url)
            cursor = conn.cursor()
            cursor.execute("SELECT name FROM sqlite_master WHERE type='table';")
            tables = cursor.fetchall()
            rows = ()
            words = []
            for table_name in tables:
                table_name = table_name[0]
                cursor.execute(f"SELECT * FROM {table_name}")
                rows = cursor.fetchall()
            for row in rows:
                for word in row:
                    if word != None:
                        words.append(word)
            self.data_all[url].extend(words)
        except FileNotFoundError:
            print("The file doesn't exist.")
        except Exception as e:
            print("An error occurred:", e)

